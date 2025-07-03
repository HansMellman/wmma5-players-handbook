# import re
# from pathlib import Path

# from docx import Document
# from docx.shared import Pt
# from fpdf import FPDF
# from jinja2 import BaseLoader, Environment

# # CONFIG
# input_path = "handbook.txt"
# word_output = "WMMA5_Handbook.docx"
# pdf_output = "WMMA5_Handbook.pdf"
# html_output = "WMMA5_Handbook.html"

# # Load cleaned input
# with open(input_path, "r", encoding="utf-8") as file:
#     lines = file.readlines()

# # Parse: title = push ... ; "string", body = mov edx ... ; "string"
# parsed_sections = []
# current_title = None
# current_body = []

# for line in lines:
#     title_match = re.search(r'push\s+\w+h\s*;\s*"(.+?)"', line)
#     body_match = re.search(r'mov edx, \w+h\s*;\s*"(.+?)"', line)

#     if title_match:
#         if current_title:
#             parsed_sections.append((current_title, " ".join(current_body)))
#         current_title = title_match.group(1).strip()
#         current_body = []
#     elif body_match and current_title:
#         current_body.append(body_match.group(1).strip())

# # Append last section
# if current_title:
#     parsed_sections.append((current_title, " ".join(current_body)))

# # --- WORD OUTPUT ---
# doc = Document()
# doc.add_heading("WMMA5 Handbook", 0)
# for header, text in parsed_sections:
#     doc.add_heading(header, level=2)
#     p = doc.add_paragraph(text)
#     p.style.font.size = Pt(11)
# doc.save(word_output)


# # --- PDF OUTPUT ---
# class PDF(FPDF):
#     def header(self):
#         self.set_font("Arial", "B", 12)
#         self.cell(0, 10, "WMMA5 Handbook", ln=True, align="C")
#         self.ln(10)

#     def chapter_title(self, title):
#         self.set_font("Arial", "B", 12)
#         self.multi_cell(0, 8, title)
#         self.ln(2)

#     def chapter_body(self, body):
#         self.set_font("Arial", "", 11)
#         self.multi_cell(0, 8, body)
#         self.ln()


# pdf = PDF()
# pdf.set_auto_page_break(auto=True, margin=15)
# pdf.add_page()
# for header, text in parsed_sections:
#     pdf.chapter_title(header)
#     pdf.chapter_body(text)
# pdf.output(pdf_output)


# # --- HTML OUTPUT ---
# def slugify(text):
#     return re.sub(r"[^\w\-]+", "-", text.strip()).lower()


# html_template = """
# <!DOCTYPE html>
# <html lang="en">
# <head>
#   <meta charset="UTF-8">
#   <title>WMMA5 Handbook</title>
#   <style>
#     body { font-family: Arial, sans-serif; padding: 2em; max-width: 900px; margin: auto; scroll-behavior: smooth; }
#     h1, h2 { color: #2c3e50; }
#     h1 { border-bottom: 2px solid #ccc; }
#     .section { margin-bottom: 2em; }
#     .toc { background: #f9f9f9; border: 1px solid #ccc; padding: 1em; margin-bottom: 2em; }
#     .toc ul { list-style-type: none; padding-left: 1em; }
#     .toc li { margin: 0.3em 0; }
#     p { word-wrap: break-word; overflow-wrap: break-word; line-height: 1.6; }
#     #top-button {
#       display: none;
#       position: fixed;
#       bottom: 30px;
#       right: 30px;
#       z-index: 99;
#       background-color: #2c3e50;
#       color: white;
#       border: none;
#       padding: 12px 16px;
#       border-radius: 6px;
#       font-size: 14px;
#       cursor: pointer;
#       opacity: 0.8;
#     }
#     #top-button:hover { background-color: #1a242f; }
#   </style>
# </head>
# <body>
#   <a id="top"></a>
#   <h1>WMMA5 Player's Handbook</h1>

#   <div class="toc">
#     <h2>Table of Contents</h2>
#     <ul>
#       {% for row in toc %}
#         <li><a href="#{{ row.id }}">{{ row.text }}</a></li>
#       {% endfor %}
#     </ul>
#   </div>

#   {% for row in rows %}
#     <div class="section">
#       <h2 id="{{ row.id }}">{{ row.HeaderText }}</h2>
#       <p>{{ row.BodyText }}</p>
#     </div>
#   {% endfor %}

#   <button onclick="scrollToTop()" id="top-button" title="Return to Top">↑ Top</button>

#   <script>
#     const topButton = document.getElementById("top-button");
#     window.onscroll = function() {
#       if (document.body.scrollTop > 400 || document.documentElement.scrollTop > 400)
#         topButton.style.display = "block";
#       else
#         topButton.style.display = "none";
#     };
#     function scrollToTop() {
#       window.scrollTo({ top: 0, behavior: 'smooth' });
#     }
#   </script>
# </body>
# </html>
# """

# env = Environment(loader=BaseLoader())
# template = env.from_string(html_template)
# rows = []
# toc = []
# for header, text in parsed_sections:
#     hid = slugify(header)
#     rows.append({"HeaderText": header, "BodyText": text, "id": hid})
#     toc.append({"id": hid, "text": header})

# rendered_html = template.render(rows=rows, toc=toc)
# Path(html_output).write_text(rendered_html, encoding="utf-8")

# print("✅ All exports complete!")

#### TAKE 2 ####

# import re
# from pathlib import Path

# from docx import Document
# from docx.shared import Pt
# from fpdf import FPDF
# from jinja2 import BaseLoader, Environment

# # CONFIG
# input_path = "handbook.txt"
# word_output = "WMMA5_Handbook.docx"
# pdf_output = "WMMA5_Handbook.pdf"
# html_output = "WMMA5_Handbook.html"

# # Load cleaned input
# with open(input_path, "r", encoding="utf-8") as file:
#     lines = file.readlines()

# # Parse: title = push ... ; "string", body = mov edx ... ; "string"
# parsed_sections = []
# current_title = None
# current_body = []

# for line in lines:
#     title_match = re.search(r'push\s+\w+h\s*;\s*"(.+?)"', line)
#     body_match = re.search(r'mov edx, \w+h\s*;\s*"(.+?)"', line)

#     if title_match:
#         if current_title:
#             parsed_sections.append((current_title, " ".join(current_body)))
#         current_title = title_match.group(1).strip()
#         current_body = []
#     elif body_match and current_title:
#         current_body.append(body_match.group(1).strip())

# # Append last section
# if current_title:
#     parsed_sections.append((current_title, " ".join(current_body)))

# # --- WORD OUTPUT ---
# doc = Document()
# doc.add_heading("WMMA5 Handbook", 0)
# for header, text in parsed_sections:
#     doc.add_heading(header, level=2)
#     p = doc.add_paragraph(text)
#     p.style.font.size = Pt(11)
# doc.save(word_output)


# # --- PDF OUTPUT ---
# class PDF(FPDF):
#     def header(self):
#         self.set_font("Arial", "B", 12)
#         self.cell(0, 10, "WMMA5 Handbook", ln=True, align="C")
#         self.ln(10)

#     def chapter_title(self, title):
#         self.set_font("Arial", "B", 12)
#         self.multi_cell(0, 8, title)
#         self.ln(2)

#     def chapter_body(self, body):
#         self.set_font("Arial", "", 11)
#         self.multi_cell(0, 8, body)
#         self.ln()


# pdf = PDF()
# pdf.set_auto_page_break(auto=True, margin=15)
# pdf.add_page()
# for header, text in parsed_sections:
#     pdf.chapter_title(header)
#     pdf.chapter_body(text)
# pdf.output(pdf_output)


# # --- HTML OUTPUT ---
# def slugify(text):
#     return re.sub(r"[^\w\-]+", "-", text.strip()).lower()


# html_template = """
# <!DOCTYPE html>
# <html lang="en">
# <head>
#   <meta charset="UTF-8">
#   <title>WMMA5 Handbook</title>
#   <style>
#     body { font-family: Arial, sans-serif; padding: 2em; max-width: 900px; margin: auto; scroll-behavior: smooth; }
#     h1, h2 { color: #2c3e50; }
#     h1 { border-bottom: 2px solid #ccc; }
#     .section { margin-bottom: 2em; }
#     .toc { background: #f9f9f9; border: 1px solid #ccc; padding: 1em; margin-bottom: 2em; }
#     .toc ul { list-style-type: none; padding-left: 1em; }
#     .toc li { margin: 0.3em 0; }
#     p { word-wrap: break-word; overflow-wrap: break-word; line-height: 1.6; }
#     #top-button {
#       display: none;
#       position: fixed;
#       bottom: 30px;
#       right: 30px;
#       z-index: 99;
#       background-color: #2c3e50;
#       color: white;
#       border: none;
#       padding: 12px 16px;
#       border-radius: 6px;
#       font-size: 14px;
#       cursor: pointer;
#       opacity: 0.8;
#     }
#     #top-button:hover { background-color: #1a242f; }
#     details summary { font-weight: bold; font-size: 1.2em; cursor: pointer; }
#   </style>
# </head>
# <body>
#   <a id="top"></a>
#   <h1>WMMA5 Player's Handbook</h1>

#   <div class="toc">
#     <h2>Table of Contents</h2>
#     <ul>
#       {% for row in toc %}
#         <li><a href="#{{ row.id }}">{{ row.text }}</a></li>
#       {% endfor %}
#     </ul>
#   </div>

#   {% for row in rows %}
#     <div class="section">
#       <details open>
#         <summary id="{{ row.id }}">{{ row.HeaderText }}</summary>
#         <p>{{ row.BodyText }}</p>
#       </details>
#     </div>
#   {% endfor %}

#   <button onclick="scrollToTop()" id="top-button" title="Return to Top">↑ Top</button>

#   <script>
#     const topButton = document.getElementById("top-button");
#     window.onscroll = function() {
#       if (document.body.scrollTop > 400 || document.documentElement.scrollTop > 400)
#         topButton.style.display = "block";
#       else
#         topButton.style.display = "none";
#     };
#     function scrollToTop() {
#       window.scrollTo({ top: 0, behavior: 'smooth' });
#     }
#   </script>
# </body>
# </html>
# """

# env = Environment(loader=BaseLoader())
# template = env.from_string(html_template)
# rows = []
# toc = []
# for header, text in parsed_sections:
#     hid = slugify(header)
#     rows.append({"HeaderText": header, "BodyText": text, "id": hid})
#     toc.append({"id": hid, "text": header})

# rendered_html = template.render(rows=rows, toc=toc)
# Path(html_output).write_text(rendered_html, encoding="utf-8")

# print("✅ All exports complete!")


import re
from pathlib import Path

from docx import Document
from docx.shared import Pt
from fpdf import FPDF
from jinja2 import BaseLoader, Environment

# CONFIG
input_path = "handbook.txt"
word_output = "WMMA5_Handbook.docx"
pdf_output = "WMMA5_Handbook.pdf"
html_output = "WMMA5_Handbook.html"

# Load cleaned input
with open(input_path, "r", encoding="utf-8") as file:
    lines = file.readlines()

# Parse: title = push ... ; "string", body = mov edx ... ; "string"
parsed_sections = []
current_title = None
current_body = []

for line in lines:
    title_match = re.search(r'push\s+\w+h\s*;\s*"(.+?)"', line)
    body_match = re.search(r'mov edx, \w+h\s*;\s*"(.+?)"', line)

    if title_match:
        if current_title:
            parsed_sections.append((current_title, " ".join(current_body)))
        current_title = title_match.group(1).strip()
        current_body = []
    elif body_match and current_title:
        current_body.append(body_match.group(1).strip())

# Append last section
if current_title:
    parsed_sections.append((current_title, " ".join(current_body)))

# --- WORD OUTPUT ---
doc = Document()
doc.add_heading("WMMA5 Handbook", 0)
for header, text in parsed_sections:
    doc.add_heading(header, level=2)
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)
doc.save(word_output)


# --- PDF OUTPUT ---
class PDF(FPDF):
    def header(self):
        self.set_font("Arial", "B", 12)
        self.cell(0, 10, "WMMA5 Handbook", ln=True, align="C")
        self.ln(10)

    def chapter_title(self, title):
        self.set_font("Arial", "B", 12)
        self.multi_cell(0, 8, title)
        self.ln(2)

    def chapter_body(self, body):
        self.set_font("Arial", "", 11)
        self.multi_cell(0, 8, body)
        self.ln()


pdf = PDF()
pdf.set_auto_page_break(auto=True, margin=15)
pdf.add_page()
for header, text in parsed_sections:
    pdf.chapter_title(header)
    pdf.chapter_body(text)
pdf.output(pdf_output)


# --- HTML OUTPUT ---
def slugify(text):
    return re.sub(r"[^\w\-]+", "-", text.strip()).lower()


html_template = """
<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <title>WMMA5 Handbook</title>
  <style>
    body {
      font-family: Arial, sans-serif;
      padding: 2em;
      max-width: 1000px;
      margin: auto;
      scroll-behavior: smooth;
      transition: background-color 0.3s, color 0.3s;
      background-color: #ffffff;
      color: #000000;
    }
    h1, h2, summary { color: #2c3e50; }
    h1 { border-bottom: 2px solid #ccc; }
    .section { margin-bottom: 2em; }

    .toc {
      background: #f9f9f9;
      border: 1px solid #ccc;
      padding: 1em;
      margin-bottom: 2em;
    }
    .toc ul { list-style-type: none; padding-left: 1em; }
    .toc li { margin: 0.3em 0; }

    p {
      word-wrap: break-word;
      overflow-wrap: break-word;
      line-height: 1.6;
    }

    a {
      color: #1a0dab;
      text-decoration: none;
    }
    a:hover {
      text-decoration: underline;
    }

    #top-button {
      display: none;
      position: fixed;
      bottom: 30px;
      right: 30px;
      z-index: 99;
      background-color: #2c3e50;
      color: white;
      border: none;
      padding: 12px 16px;
      border-radius: 6px;
      font-size: 14px;
      cursor: pointer;
      opacity: 0.8;
    }
    #top-button:hover { background-color: #1a242f; }

    details summary {
      font-weight: bold;
      font-size: 1.2em;
      cursor: pointer;
    }

    #controls {
      display: flex;
      justify-content: space-between;
      margin-bottom: 1.5em;
    }
    #search-box {
      width: 60%;
      padding: 8px;
      font-size: 1em;
    }
    #theme-toggle {
      padding: 8px 12px;
      font-size: 1em;
      cursor: pointer;
    }

    /* Dark Mode Styles */
    .dark-mode {
      background-color: #1a1a1a;
      color: #f0f0f0;
    }
    .dark-mode .toc {
      background: #2c2c2c;
      border-color: #555;
    }
    .dark-mode summary,
    .dark-mode h1,
    .dark-mode h2 {
      color: #f0f0f0;
    }
    .dark-mode a {
      color: #4ea1ff;
    }
  </style>
</head>
<body>
  <a id="top"></a>
  <h1>WMMA5 Player's Handbook</h1>

  <div id="controls">
    <input type="text" id="search-box" placeholder="Search for a section...">
    <button id="theme-toggle">Dark/Light</button>
  </div>

  <div class="toc">
    <h2>Table of Contents</h2>
    <ul>
      {% for row in toc %}
        <li><a href="#{{ row.id }}">{{ row.text }}</a></li>
      {% endfor %}
    </ul>
  </div>

  {% for row in rows %}
    <div class="section">
      <details open>
        <summary id="{{ row.id }}">{{ row.HeaderText }}</summary>
        <p>{{ row.BodyText }}</p>
      </details>
    </div>
  {% endfor %}

  <button onclick="scrollToTop()" id="top-button" title="Return to Top">↑ Top</button>

<script>
  const topButton = document.getElementById("top-button");
  const themeToggle = document.getElementById("theme-toggle");
  const searchBox = document.getElementById("search-box");

  window.onscroll = function() {
    topButton.style.display = (document.body.scrollTop > 400 || document.documentElement.scrollTop > 400) ? "block" : "none";
  };

  function scrollToTop() {
    window.scrollTo({ top: 0, behavior: 'smooth' });
  }

  function setTheme(theme) {
    document.body.classList.toggle("dark-mode", theme === "dark");
    localStorage.setItem("theme", theme);
  }

  themeToggle.addEventListener("click", function () {
    const newTheme = document.body.classList.contains("dark-mode") ? "light" : "dark";
    setTheme(newTheme);
  });

  document.addEventListener("DOMContentLoaded", () => {
    const savedTheme = localStorage.getItem("theme") || "light";
    setTheme(savedTheme);
  });

  searchBox.addEventListener("input", function () {
    const query = this.value.toLowerCase();

    // Filter section blocks
    document.querySelectorAll(".section").forEach(section => {
      const text = section.innerText.toLowerCase();
      section.style.display = text.includes(query) ? "block" : "none";
    });

    // Filter Table of Contents links
    document.querySelectorAll(".toc li").forEach(li => {
      const text = li.innerText.toLowerCase();
      li.style.display = text.includes(query) ? "list-item" : "none";
    });
  });

  // Optional: scroll to first match on Enter
  searchBox.addEventListener("keydown", function (event) {
    if (event.key === "Enter") {
      const query = this.value.toLowerCase();
      const first = Array.from(document.querySelectorAll(".section")).find(section =>
        section.innerText.toLowerCase().includes(query)
      );
      if (first) {
        first.scrollIntoView({ behavior: 'smooth' });
      }
    }
  });
</script>
</body>
</html>
"""

env = Environment(loader=BaseLoader())
template = env.from_string(html_template)
rows = []
toc = []
for header, text in parsed_sections:
    hid = slugify(header)
    rows.append({"HeaderText": header, "BodyText": text, "id": hid})
    toc.append({"id": hid, "text": header})

rendered_html = template.render(rows=rows, toc=toc)
Path(html_output).write_text(rendered_html, encoding="utf-8")

print("✅ All exports complete!")
